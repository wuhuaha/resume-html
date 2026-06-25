from __future__ import annotations

from pathlib import Path
from tempfile import NamedTemporaryFile
from zipfile import ZipFile
from xml.etree import ElementTree as ET

from docx import Document
from fastapi import UploadFile
from pypdf import PdfReader


async def upload_to_markdown(file: UploadFile) -> str:
    suffix = Path(file.filename or "").suffix.lower()
    content = await file.read()

    with NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        tmp.write(content)
        tmp_path = Path(tmp.name)

    try:
        if suffix == ".docx":
            return docx_to_markdown(tmp_path)
        if suffix == ".pdf":
            return pdf_to_markdown(tmp_path)
        if suffix in {".md", ".markdown", ".txt"}:
            return content.decode("utf-8", errors="ignore")
        raise ValueError("仅支持 .docx、.pdf、.md、.txt 文件")
    finally:
        tmp_path.unlink(missing_ok=True)


def docx_to_markdown(path: Path) -> str:
    document = Document(path)
    lines: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style_name = paragraph.style.name.lower()
        if "heading 1" in style_name:
            lines.append(f"# {text}")
        elif "heading 2" in style_name:
            lines.append(f"## {text}")
        elif "heading 3" in style_name:
            lines.append(f"### {text}")
        else:
            lines.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " / ") for cell in row.cells]
            if any(cells):
                lines.append(" | ".join(cells))

    markdown = "\n\n".join(lines).strip()
    if markdown:
        return markdown

    return docx_xml_text_to_markdown(path)


def docx_xml_text_to_markdown(path: Path) -> str:
    texts: list[str] = []
    seen: set[str] = set()
    with ZipFile(path) as archive:
        for name in archive.namelist():
            if not name.startswith("word/") or not name.endswith(".xml"):
                continue
            try:
                root = ET.fromstring(archive.read(name))
            except ET.ParseError:
                continue
            for paragraph in root.iter("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p"):
                parts = [
                    node.text
                    for node in paragraph.iter("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t")
                    if node.text
                ]
                text = "".join(parts).strip()
                if not text or text in seen:
                    continue
                seen.add(text)
                texts.append(text)

    return paragraphs_to_markdown(texts)


def paragraphs_to_markdown(paragraphs: list[str]) -> str:
    section_titles = {
        "个人简介",
        "基本信息",
        "核心能力",
        "技术栈 Education",
        "工作经历 Work experience",
        "奖项荣誉 Awards honor",
        "技能证书 Skill certificate",
        "自我评价 Self evaluation",
        "项目经验 Project experience",
    }
    lines: list[str] = []
    for text in paragraphs:
        if text in section_titles:
            title = text.split()[0]
            lines.append(f"## {title}")
        elif text.startswith("【项目名称】"):
            lines.append(f"### {text.replace('【项目名称】', '').strip()}")
        elif text.startswith("【项目职位】") or text.startswith("【项目内容】"):
            lines.append(text)
        elif text.endswith("；") or text.endswith("。"):
            lines.append(f"- {text}")
        else:
            lines.append(text)
    return "\n\n".join(lines).strip()


def pdf_to_markdown(path: Path) -> str:
    reader = PdfReader(str(path))
    pages = []
    for index, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        if text.strip():
            pages.append(f"## Page {index}\n\n{text.strip()}")
    return "\n\n".join(pages).strip()
